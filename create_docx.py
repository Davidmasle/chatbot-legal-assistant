from docx import Document

doc = Document()
doc.add_paragraph("Test docx file for adada  Legal Assistant")
doc.add_paragraph("Example text should be divided in chunks addadasd a dasdasdasd dadad")
doc.save("sample.docx")

print("DOCX created: sample.docx")
